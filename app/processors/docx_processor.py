import docx
from typing import Dict, Any
from app.processors.base_processor import BaseProcessor
from app.utils.logger import logger

class DOCXProcessor(BaseProcessor):
    def extract_text(self, file_path: str) -> str:
        """Extrai texto de um arquivo DOCX."""
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            logger.error(f"Erro ao extrair texto do DOCX {file_path}: {e}")
        return text

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extrai metadados do DOCX."""
        metadata = {}
        try:
            doc = docx.Document(file_path)
            prop = doc.core_properties
            metadata = {
                'author': prop.author,
                'category': prop.category,
                'comments': prop.comments,
                'content_status': prop.content_status,
                'created': str(prop.created),
                'identifier': prop.identifier,
                'keywords': prop.keywords,
                'last_modified_by': prop.last_modified_by,
                'language': prop.language,
                'modified': str(prop.modified),
                'subject': prop.subject,
                'title': prop.title,
                'version': prop.version
            }
        except Exception as e:
            logger.error(f"Erro ao extrair metadados do DOCX {file_path}: {e}")
        return metadata
